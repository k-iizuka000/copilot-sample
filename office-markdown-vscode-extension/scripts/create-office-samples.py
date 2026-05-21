#!/usr/bin/env python3

from __future__ import annotations

import os
import shutil
import subprocess
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.drawing.image import Image as XlsxImage
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "samples" / "windows-office-fixtures"
ASSET_DIR = OUTPUT_DIR / "source-assets"

WORD_OUT = OUTPUT_DIR / "word-image-object-sample.docx"
EXCEL_OUT = OUTPUT_DIR / "excel-image-object-sample.xlsx"
POWERPOINT_OUT = OUTPUT_DIR / "powerpoint-image-object-sample.pptx"
IMAGE_OUT = ASSET_DIR / "office-markdown-sample-visual.png"

REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
XDR_NS = "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

ET.register_namespace("", REL_NS)
ET.register_namespace("ct", CT_NS)
ET.register_namespace("xdr", XDR_NS)
ET.register_namespace("a", A_NS)


def q(ns: str, tag: str) -> str:
    return f"{{{ns}}}{tag}"


def rewrite_zip(zip_path: Path, replacements: dict[str, bytes | str], additions: dict[str, bytes]) -> None:
    temp_path = zip_path.with_suffix(zip_path.suffix + ".tmp")
    skip = set(replacements) | set(additions)
    with zipfile.ZipFile(zip_path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for info in source.infolist():
            if info.filename in skip:
                continue
            target.writestr(info, source.read(info.filename))
        for name, data in replacements.items():
            target.writestr(name, data.encode("utf-8") if isinstance(data, str) else data)
        for name, data in additions.items():
            target.writestr(name, data)
    temp_path.replace(zip_path)


def zip_read_optional(zip_path: Path, entry: str) -> bytes | None:
    with zipfile.ZipFile(zip_path, "r") as archive:
        try:
            return archive.read(entry)
        except KeyError:
            return None


def ensure_bin_content_type(xml_bytes: bytes) -> bytes:
    root = ET.fromstring(xml_bytes)
    for child in root:
        if child.tag.endswith("Default") and child.attrib.get("Extension") == "bin":
            return ET.tostring(root, encoding="utf-8", xml_declaration=True)
    ET.SubElement(root, q(CT_NS, "Default"), Extension="bin", ContentType="application/octet-stream")
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def append_relationship(
    rels_bytes: bytes | None,
    rel_id: str,
    rel_type: str,
    target: str,
) -> bytes:
    if rels_bytes:
        root = ET.fromstring(rels_bytes)
    else:
        root = ET.Element(q(REL_NS, "Relationships"))
    for child in root:
        if child.attrib.get("Id") == rel_id:
            return ET.tostring(root, encoding="utf-8", xml_declaration=True)
    ET.SubElement(root, q(REL_NS, "Relationship"), Id=rel_id, Type=rel_type, Target=target)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def make_visual_asset() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (960, 520), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 34)
        body_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 20)
        small_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 16)
    except OSError:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    draw.rectangle((0, 0, 960, 520), fill="#F8FAFC")
    draw.rectangle((36, 36, 924, 484), outline="#CBD5E1", width=3)
    draw.text((64, 58), "Office Markdown sample visual", fill="#0F172A", font=title_font)
    draw.text((66, 108), "Embedded into Word, Excel, and PowerPoint fixtures", fill="#475569", font=body_font)

    boxes = [
        ("Text", "#2563EB", 84, 190),
        ("Image", "#0F766E", 374, 190),
        ("Object", "#B45309", 664, 190),
    ]
    for label, color, x, y in boxes:
        draw.rounded_rectangle((x, y, x + 210, y + 112), radius=18, fill=color)
        draw.text((x + 58, y + 36), label, fill="#FFFFFF", font=body_font)

    draw.line((294, 246, 374, 246), fill="#64748B", width=5)
    draw.polygon([(374, 246), (352, 234), (352, 258)], fill="#64748B")
    draw.line((584, 246, 664, 246), fill="#64748B", width=5)
    draw.polygon([(664, 246), (642, 234), (642, 258)], fill="#64748B")

    draw.rectangle((84, 358, 876, 414), fill="#FFFFFF", outline="#CBD5E1", width=2)
    draw.text((108, 374), "Purpose: verify extraction of text, tables, pictures, drawings, charts, and objects.", fill="#334155", font=small_font)
    image.save(IMAGE_OUT)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    props.append(style)
    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_word_sample() -> None:
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.color.rgb = RGBColor(31, 41, 55)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Office Markdown Fixture - Word")

    lead = document.add_paragraph()
    lead.add_run("Windows-focused DOCX sample with ")
    add_hyperlink(lead, "external hyperlink", "https://example.com/office-markdown-fixture")
    lead.add_run(", table content, an inline image, and a synthetic embedded object relationship.")

    jp = "\u65e5\u672c\u8a9e\u30c6\u30ad\u30b9\u30c8: Windows \u691c\u8a3c\u7528\u306e Unicode \u884c\u3067\u3059\u3002"
    document.add_paragraph(jp)

    document.add_heading("Image and caption", level=1)
    document.add_picture(str(IMAGE_OUT), width=Inches(5.8))
    caption = document.add_paragraph("Figure 1. Embedded PNG image used by all three Office samples.")
    caption.style = "Caption"

    document.add_heading("Conversion coverage table", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Element"
    header[1].text = "Expected Markdown signal"
    header[2].text = "Notes"
    rows = [
        ("Heading", "# Image and caption", "Word heading style"),
        ("Hyperlink", "[external hyperlink](...)", "External relationship"),
        ("Image", "doc-image-001.png", "Inline drawing relationship"),
        ("Object", "doc-object-001.bin", "Synthetic OLE relationship"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value

    document.add_heading("Visible object placeholder", level=1)
    object_box = document.add_table(rows=1, cols=1)
    object_box.style = "Table Grid"
    cell = object_box.rows[0].cells[0]
    cell.text = "Embedded object placeholder: sample-object.bin is included in word/embeddings."
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "E0F2FE")
    cell._tc.get_or_add_tcPr().append(shading)

    document.save(WORD_OUT)

    rels_path = "word/_rels/document.xml.rels"
    types_path = "[Content_Types].xml"
    replacements = {
        rels_path: append_relationship(
            zip_read_optional(WORD_OUT, rels_path),
            "rIdOfficeObject1",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject",
            "embeddings/sample-object.bin",
        ),
        types_path: ensure_bin_content_type(zip_read_optional(WORD_OUT, types_path) or b""),
    }
    rewrite_zip(
        WORD_OUT,
        replacements,
        {
            "word/embeddings/sample-object.bin": b"Synthetic embedded object payload for Office Markdown DOCX fixture.\n",
        },
    )


def create_excel_sample() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Summary"
    sheet.sheet_view.showGridLines = False
    sheet["A1"] = "Office Markdown Fixture - Excel"
    sheet.merge_cells("A1:D1")
    sheet["A1"].font = Font(name="Aptos Display", size=14, bold=True, color="1F2937")
    sheet["A2"] = "Formulas, chart, image, drawing textbox, hidden sheet, and embedded object relationship."
    sheet["A2"].font = Font(name="Aptos", size=9, color="64748B")
    sheet["A3"] = "\u65e5\u672c\u8a9e\u30c6\u30ad\u30b9\u30c8: Windows \u691c\u8a3c\u7528\u3002"

    header_row = 14
    data_start_row = header_row + 1
    formula_row = header_row + 5
    headers = ["Category", "Count", "Coverage", "Expected signal"]
    for col, header in enumerate(headers, start=1):
        cell = sheet.cell(row=header_row, column=col, value=header)
        cell.fill = PatternFill("solid", fgColor="1F2937")
        cell.font = Font(name="Aptos", size=10, bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center")

    rows = [
        ("Text blocks", 8, 0.92, "Markdown table cells"),
        ("Images", 1, 1.0, "sheet image asset"),
        ("Objects", 1, 1.0, "sheet object asset"),
    ]
    for row_index, row_values in enumerate(rows, start=data_start_row):
        for col_index, value in enumerate(row_values, start=1):
            cell = sheet.cell(row=row_index, column=col_index, value=value)
            cell.font = Font(name="Aptos", size=10, color="111827")
            cell.alignment = Alignment(vertical="center")
            if col_index == 3:
                cell.number_format = "0%"

    sheet.cell(row=formula_row, column=1, value="Total")
    sheet.cell(row=formula_row, column=2, value=f"=SUM(B{data_start_row}:B{data_start_row + 2})")
    sheet.cell(row=formula_row, column=3, value=f"=AVERAGE(C{data_start_row}:C{data_start_row + 2})")
    sheet.cell(row=formula_row, column=3).number_format = "0%"
    sheet.cell(row=formula_row, column=4, value="Formula cells should appear in manifest")
    for cell in sheet[formula_row]:
        cell.font = Font(name="Aptos", size=10, bold=True, color="111827")
        cell.fill = PatternFill("solid", fgColor="E0F2FE")

    for width_index, width in enumerate([16, 9, 10, 26], start=1):
        sheet.column_dimensions[get_column_letter(width_index)].width = width
    for row in range(header_row, formula_row + 1):
        sheet.row_dimensions[row].height = 19
    sheet.sheet_view.zoomScale = 80
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 1
    sheet.print_area = "A1:J24"

    img = XlsxImage(str(IMAGE_OUT))
    img.width = 300
    img.height = 162
    sheet.add_image(img, "A5")

    chart = BarChart()
    chart.title = "Fixture coverage"
    chart.y_axis.title = "Count"
    chart.x_axis.title = "Category"
    data = Reference(sheet, min_col=2, min_row=header_row, max_row=data_start_row + 2)
    cats = Reference(sheet, min_col=1, min_row=data_start_row, max_row=data_start_row + 2)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 7
    chart.width = 11
    sheet.add_chart(chart, "E5")

    hidden = workbook.create_sheet("Hidden Notes")
    hidden["A1"] = "This hidden sheet should be recorded or skipped depending on extension settings."
    hidden.sheet_state = "hidden"

    workbook.save(EXCEL_OUT)
    patch_excel_ooxml()


def set_cached_formula_values(sheet_xml: bytes) -> bytes:
    root = ET.fromstring(sheet_xml)
    ns = {"ws": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    cached_values = {"B19": "10", "C19": "0.9733333333"}
    for ref, value in cached_values.items():
        cell = root.find(f".//ws:c[@r='{ref}']", ns)
        if cell is None:
            continue
        value_node = cell.find("ws:v", ns)
        if value_node is None:
            value_node = ET.SubElement(cell, q(ns["ws"], "v"))
        value_node.text = value
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def add_excel_textbox(drawing_xml: bytes) -> bytes:
    root = ET.fromstring(drawing_xml)
    anchor = ET.Element(q(XDR_NS, "oneCellAnchor"))
    from_node = ET.SubElement(anchor, q(XDR_NS, "from"))
    for tag, value in [("col", "4"), ("colOff", "0"), ("row", "9"), ("rowOff", "0")]:
        ET.SubElement(from_node, q(XDR_NS, tag)).text = value
    ET.SubElement(anchor, q(XDR_NS, "ext"), cx="3600000", cy="780000")

    shape = ET.SubElement(anchor, q(XDR_NS, "sp"), macro="", textlink="")
    nv = ET.SubElement(shape, q(XDR_NS, "nvSpPr"))
    ET.SubElement(nv, q(XDR_NS, "cNvPr"), id="1025", name="Office Markdown TextBox")
    ET.SubElement(nv, q(XDR_NS, "cNvSpPr"), txBox="1")
    sp_pr = ET.SubElement(shape, q(XDR_NS, "spPr"))
    solid = ET.SubElement(sp_pr, q(A_NS, "solidFill"))
    ET.SubElement(solid, q(A_NS, "srgbClr"), val="E0F2FE")
    line = ET.SubElement(sp_pr, q(A_NS, "ln"))
    line_fill = ET.SubElement(line, q(A_NS, "solidFill"))
    ET.SubElement(line_fill, q(A_NS, "srgbClr"), val="0284C7")
    geom = ET.SubElement(sp_pr, q(A_NS, "prstGeom"), prst="roundRect")
    ET.SubElement(geom, q(A_NS, "avLst"))

    body = ET.SubElement(shape, q(XDR_NS, "txBody"))
    ET.SubElement(body, q(A_NS, "bodyPr"), wrap="square", rtlCol="0")
    ET.SubElement(body, q(A_NS, "lstStyle"))
    para = ET.SubElement(body, q(A_NS, "p"))
    run = ET.SubElement(para, q(A_NS, "r"))
    run_props = ET.SubElement(run, q(A_NS, "rPr"), lang="en-US", sz="1100")
    run_fill = ET.SubElement(run_props, q(A_NS, "solidFill"))
    ET.SubElement(run_fill, q(A_NS, "srgbClr"), val="0F172A")
    ET.SubElement(run, q(A_NS, "t")).text = "Important textbox: drawing text should be captured."
    ET.SubElement(anchor, q(XDR_NS, "clientData"))
    root.append(anchor)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def patch_excel_ooxml() -> None:
    sheet_rels = "xl/worksheets/_rels/sheet1.xml.rels"
    drawing_path = "xl/drawings/drawing1.xml"
    replacements = {
        "[Content_Types].xml": ensure_bin_content_type(zip_read_optional(EXCEL_OUT, "[Content_Types].xml") or b""),
        "xl/worksheets/sheet1.xml": set_cached_formula_values(
            zip_read_optional(EXCEL_OUT, "xl/worksheets/sheet1.xml") or b""
        ),
        sheet_rels: append_relationship(
            zip_read_optional(EXCEL_OUT, sheet_rels),
            "rIdOfficeObject1",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject",
            "../embeddings/sample-object.bin",
        ),
        drawing_path: add_excel_textbox(zip_read_optional(EXCEL_OUT, drawing_path) or b""),
    }
    rewrite_zip(
        EXCEL_OUT,
        replacements,
        {
            "xl/embeddings/sample-object.bin": b"Synthetic embedded object payload for Office Markdown XLSX fixture.\n",
        },
    )


def create_powerpoint_sample() -> None:
    node = os.environ.get(
        "CODEX_NODE",
        str(
            Path.home()
            / ".cache"
            / "codex-runtimes"
            / "codex-primary-runtime"
            / "dependencies"
            / "node"
            / "bin"
            / "node"
        ),
    )
    if not Path(node).exists():
        node = shutil.which("node") or node
    script = ROOT / "scripts" / "create-powerpoint-sample.cjs"
    env = os.environ.copy()
    env.setdefault(
        "CODEX_NODE_MODULES",
        str(
            Path.home()
            / ".cache"
            / "codex-runtimes"
            / "codex-primary-runtime"
            / "dependencies"
            / "node"
            / "node_modules"
        ),
    )
    subprocess.run(
        [
            node,
            str(script),
            "--out",
            str(POWERPOINT_OUT),
            "--image",
            str(IMAGE_OUT),
        ],
        check=True,
        env=env,
    )


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    make_visual_asset()
    create_word_sample()
    create_excel_sample()
    create_powerpoint_sample()
    print(f"Created {WORD_OUT}")
    print(f"Created {EXCEL_OUT}")
    print(f"Created {POWERPOINT_OUT}")


if __name__ == "__main__":
    main()
